"""
File handler module for reading and writing PDF and DOCX files.
"""
import os
import io
import logging
from typing import Tuple, Optional
import PyPDF2
import docx
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from docx import Document

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

def open_file_dialog(file_types: Tuple[Tuple[str, str], ...]) -> Optional[str]:
    """
    Open a file dialog to select a file.
    
    Args:
        file_types: Tuple of file type descriptions and extensions (e.g., (("PDF Files", "*.pdf"),)).
        
    Returns:
        The selected file path or None if no file was selected.
    """
    try:
        import tkinter as tk
        from tkinter import filedialog
        
        # Create and configure the root window
        root = tk.Tk()
        root.withdraw()  # Hide the main window
        root.attributes('-topmost', True)  # Make sure the dialog stays on top
        
        logger.info("Opening file selection dialog")
        file_path = filedialog.askopenfilename(
            title="Select a file",
            filetypes=file_types,
            parent=root
        )
        
        # Destroy the root window
        root.destroy()
        
        if file_path:
            logger.info(f"Selected file: {file_path}")
            return file_path
        else:
            logger.info("No file selected")
            return None
            
    except Exception as e:
        logger.error(f"Error opening file dialog: {str(e)}")
        return None

def save_file_dialog(default_filename: str, file_types: Tuple[Tuple[str, str], ...]) -> Optional[str]:
    """
    Open a save file dialog to select a save path.
    
    Args:
        default_filename: Default filename to suggest.
        file_types: Tuple of file type descriptions and extensions (e.g., (("PDF Files", "*.pdf"),)).
        
    Returns:
        The selected save path or None if no path was selected.
    """
    try:
        import tkinter as tk
        from tkinter import filedialog
        
        # Create and configure the root window
        root = tk.Tk()
        root.withdraw()  # Hide the main window
        root.attributes('-topmost', True)  # Make sure the dialog stays on top
        
        # Get the file extension from the first file type
        default_ext = file_types[0][1].replace("*", "")
        
        logger.info(f"Opening save dialog with default filename: {default_filename}")
        file_path = filedialog.asksaveasfilename(
            title="Save As",
            defaultextension=default_ext,
            initialfile=default_filename,
            filetypes=file_types,
            parent=root
        )
        
        # Destroy the root window
        root.destroy()
        
        if file_path:
            logger.info(f"Selected save path: {file_path}")
            return file_path
        else:
            logger.info("No save path selected")
            return None
            
    except Exception as e:
        logger.error(f"Error opening save file dialog: {str(e)}")
        return None

def read_pdf(file_path: str) -> Optional[str]:
    """
    Read text from a PDF file.
    
    Args:
        file_path: Path to the PDF file.
        
    Returns:
        The extracted text or None if an error occurred.
    """
    try:
        text = ""
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text += page.extract_text()
        
        if not text.strip():
            logger.warning(f"No text extracted from PDF file: {file_path}")
        
        return text
    except Exception as e:
        logger.error(f"Error reading PDF file {file_path}: {str(e)}")
        return None

def read_docx(file_path: str) -> Optional[str]:
    """
    Read text from a DOCX file.
    
    Args:
        file_path: Path to the DOCX file.
        
    Returns:
        The extracted text or None if an error occurred.
    """
    try:
        doc = Document(file_path)
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        
        if not text.strip():
            logger.warning(f"No text extracted from DOCX file: {file_path}")
        
        return text
    except Exception as e:
        logger.error(f"Error reading DOCX file {file_path}: {str(e)}")
        return None

def save_file(text: str, file_path: str) -> bool:
    """
    Save text to a file.
    
    Args:
        text: Text to save.
        file_path: Path to save the file.
        
    Returns:
        True if successful, False otherwise.
    """
    try:
        if not text:
            logger.error("No text content to save.")
            return False

        _, ext = os.path.splitext(file_path)
        ext = ext.lower()
        
        logger.info(f"Attempting to save file with extension: {ext}")
        
        if ext == '.pdf':
            return save_as_pdf(text, file_path)
        elif ext == '.docx':
            return save_as_docx(text, file_path)
        else:
            logger.error(f"Unsupported file format: {ext}")
            return False
    except Exception as e:
        logger.error(f"Error saving file {file_path}: {str(e)}")
        return False

def save_as_pdf(text: str, file_path: str) -> bool:
    """
    Save text as a PDF file.
    
    Args:
        text: Text to save.
        file_path: Path to save the PDF file.
        
    Returns:
        True if successful, False otherwise.
    """
    try:
        logger.info("Creating PDF document...")
        # Create a PDF document
        packet = io.BytesIO()
        c = canvas.Canvas(packet, pagesize=letter)
        
        # Set font and size
        c.setFont("Helvetica", 11)  # Increased font size
        
        # Calculate page dimensions
        width, height = letter
        margin = 72  # 1 inch margin
        line_height = 14  # Increased line height
        
        # Split text into paragraphs
        paragraphs = text.split('\n\n')  # Split by double newline for paragraphs
        
        # Add text to the PDF
        y = height - margin
        for paragraph in paragraphs:
            # Split paragraph into lines that fit the page width
            lines = []
            words = paragraph.split()
            current_line = []
            
            for word in words:
                # Check if adding this word would exceed the line width
                test_line = ' '.join(current_line + [word])
                if c.stringWidth(test_line, "Helvetica", 11) < (width - 2 * margin):
                    current_line.append(word)
                else:
                    lines.append(' '.join(current_line))
                    current_line = [word]
            
            if current_line:
                lines.append(' '.join(current_line))
            
            # Add each line to the PDF
            for line in lines:
                # Add a new page if we've reached the bottom margin
                if y < margin:
                    c.showPage()
                    c.setFont("Helvetica", 11)
                    y = height - margin
                
                # Write the line
                c.drawString(margin, y, line)
                y -= line_height
            
            # Add extra space after paragraph
            y -= line_height / 2
        
        logger.info("Saving PDF document...")
        # Save the PDF
        c.save()
        
        # Write the PDF to file
        with open(file_path, 'wb') as file:
            file.write(packet.getvalue())
        
        logger.info(f"PDF saved successfully: {file_path}")
        return True
    except Exception as e:
        logger.error(f"Error saving PDF file {file_path}: {str(e)}")
        return False

def save_as_docx(text: str, file_path: str) -> bool:
    """
    Save text as a DOCX file.
    
    Args:
        text: Text to save.
        file_path: Path to save the DOCX file.
        
    Returns:
        True if successful, False otherwise.
    """
    try:
        logger.info("Creating DOCX document...")
        # Create a new Document
        doc = Document()
        
        # Split text into paragraphs
        paragraphs = text.split('\n')
        
        # Add each paragraph to the document
        for paragraph in paragraphs:
            if paragraph.strip():  # Skip empty paragraphs
                doc.add_paragraph(paragraph)
        
        logger.info("Saving DOCX document...")
        # Save the document
        doc.save(file_path)
        
        logger.info(f"DOCX saved successfully: {file_path}")
        return True
    except Exception as e:
        logger.error(f"Error saving DOCX file {file_path}: {str(e)}")
        return False

def read_file(file_path: str) -> Optional[Tuple[str, str]]:
    """
    Read a file and extract its text.
    
    Args:
        file_path: Path to the file.
        
    Returns:
        A tuple containing the extracted text and the file extension, or None if an error occurred.
    """
    try:
        _, ext = os.path.splitext(file_path)
        ext = ext.lower()
        
        if ext == '.pdf':
            text = read_pdf(file_path)
            return (text, ext) if text else None
        elif ext == '.docx':
            text = read_docx(file_path)
            return (text, ext) if text else None
        else:
            logger.error(f"Unsupported file format: {ext}")
            return None
    except Exception as e:
        logger.error(f"Error reading file {file_path}: {str(e)}")
        return None 